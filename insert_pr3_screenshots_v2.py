"""
insert_pr3_screenshots_v2.py — Вставляет скриншоты в ПР3 (обновлённое имя)
"""

import os
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC = r'C:\Users\大猪鼻\Desktop\godss\Практическая работа №3.docx'
DIR = r'C:\Users\大猪鼻\Desktop\godss'

FIGURES = {
    1: '截图1.png',
    2: '截图2.png',
    3: '截图3.png',
    4: '截图4.png',
}

doc = Document(DOC)
WIDTH_CM = 14
inserted = 0

for p in doc.paragraphs:
    text = p.text.strip()
    if 'Рисунок' in text and '–' in text:
        # Находим рисунки типа "Рисунок 2 – «Создание нового запроса»"
        for num, filename in FIGURES.items():
            if str(num) in text:
                img_path = os.path.join(DIR, filename)
                if os.path.exists(img_path):
                    # Добавляем изображение ПЕРЕД этим параграфом
                    new_p = doc.add_paragraph()
                    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = new_p.add_run()
                    run.add_picture(img_path, width=Cm(WIDTH_CM))
                    # Вставляем перед текущим параграфом
                    p._element.addprevious(new_p._element)
                    inserted += 1
                    print(f'  ✅ Рисунок {num} вставлен')
                    break

print(f'\n✅ Вставлено {inserted} из {len(FIGURES)} скриншотов')
doc.save(DOC)
print(f'✅ Документ сохранён: {DOC}')
