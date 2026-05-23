"""
practical_work2.py — Генератор документа ПР №2
===============================================
7 тест-кейсов для туристической ИС.
"""

import os
from docx import Document
from docx.shared import Pt, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PAGE_W = 7560310
PAGE_H = 10692130
MARGIN_TOP = 720090
MARGIN_BOT = 720090
MARGIN_LEFT = 1080135
MARGIN_RIGHT = 540385
SIZE_TITLE = 28
SIZE_SUBTITLE = 18
SIZE_HEADER = 16
SIZE_BODY = 14
SIZE_SMALL = 12
FONT_NAME = 'Times New Roman'
CENTER = WD_ALIGN_PARAGRAPH.CENTER
LEFT = WD_ALIGN_PARAGRAPH.LEFT
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY

def set_font(run, name=FONT_NAME, size=SIZE_BODY, bold=None):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def mkp(doc, text, size=SIZE_BODY, bold=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, name=FONT_NAME, size=size, bold=bold)
    return p

def mkp_multi(doc, parts, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    for text, size, bold in parts:
        run = p.add_run(text)
        set_font(run, name=FONT_NAME, size=size, bold=bold)
    return p

def blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()

def add_test_case_table(doc, tc_id, title, preconditions, steps, expected):
    """Добавляет таблицу тест-кейса."""
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    data = [
        ('№', tc_id),
        ('Название', title),
        ('Предусловия', preconditions),
        ('Последовательность действий', steps),
        ('Ожидаемый результат', expected),
    ]
    
    for i, (label, value) in enumerate(data):
        cell0 = table.rows[i].cells[0]
        cell1 = table.rows[i].cells[1]
        
        # Label cell
        cell0.text = ''
        run0 = cell0.paragraphs[0].add_run(label)
        set_font(run0, size=SIZE_SMALL, bold=True)
        
        # Value cell
        cell1.text = ''
        run1 = cell1.paragraphs[0].add_run(value)
        set_font(run1, size=SIZE_SMALL, bold=False)
        
        # Set widths
        cell0.width = Cm(3.5)
        cell1.width = Cm(12)

def create():
    doc = Document()
    s = doc.sections[0]
    s.page_width = Emu(PAGE_W)
    s.page_height = Emu(PAGE_H)
    s.top_margin = Emu(MARGIN_TOP)
    s.bottom_margin = Emu(MARGIN_BOT)
    s.left_margin = Emu(MARGIN_LEFT)
    s.right_margin = Emu(MARGIN_RIGHT)

    # Шапка
    mkp_multi(doc, [
        ('Министерство науки и высшего образования Российской Федерации '
         'Федеральное государственное бюджетное образовательное учреждение '
         'высшего образования',
         SIZE_BODY, False),
        ('«Владимирский государственный университет имени Александра '
         'Григорьевича и Николая Григорьевича Столетовых» (ВлГУ)',
         SIZE_BODY, True),
    ], align=CENTER)
    mkp(doc, 'Кафедра информационных систем и программной инженерии',
        size=SIZE_BODY, bold=False, align=CENTER)
    blank(doc, 3)

    mkp(doc, 'Практическая работа № 2', size=SIZE_TITLE, bold=True, align=CENTER)
    blank(doc, 1)
    mkp(doc, 'по дисциплине «Тестирование программного обеспечения»',
        size=SIZE_SUBTITLE, bold=True, align=CENTER)
    mkp(doc, 'Тема работы «Оформление тест-кейсов»',
        size=SIZE_SUBTITLE, bold=True, align=CENTER)
    blank(doc, 3)

    mkp(doc, 'Выполнил:', size=SIZE_BODY, bold=False, align=LEFT)
    mkp(doc, 'студент гр. ПРИк-223', size=SIZE_BODY, bold=False, align=LEFT)
    mkp(doc, 'Чэнь Гэ (Игорь)', size=SIZE_BODY, bold=False, align=LEFT)
    blank(doc, 1)
    mkp(doc, 'Принял:', size=SIZE_BODY, bold=False, align=LEFT)
    mkp(doc, 'Хлызова В. Г.', size=SIZE_BODY, bold=False, align=LEFT)
    blank(doc, 2)
    mkp(doc, 'Владимир, 2026', size=SIZE_BODY, bold=False, align=CENTER)
    doc.add_page_break()

    mkp(doc, 'Тест-кейсы для туристической информационной системы',
        size=SIZE_HEADER, bold=True, align=JUSTIFY)
    blank(doc, 1)

    test_cases = [
        (
            'TC-001',
            'Регистрация нового пользователя',
            'Пользователь открыл страницу регистрации. Все поля формы пусты.',
            '1. Ввести корректный email (например, user@test.ru)\n'
            '2. Ввести пароль (не менее 6 символов)\n'
            '3. Подтвердить пароль\n'
            '4. Ввести имя пользователя\n'
            '5. Нажать кнопку «Зарегистрироваться»',
            'Пользователь успешно зарегистрирован. '
            'Система перенаправляет на страницу входа. '
            'В базе данных создана новая запись пользователя.'
        ),
        (
            'TC-002',
            'Авторизация пользователя',
            'Пользователь зарегистрирован. Открыта страница входа.',
            '1. Ввести email зарегистрированного пользователя\n'
            '2. Ввести корректный пароль\n'
            '3. Нажать кнопку «Войти»',
            'Пользователь успешно авторизован. '
            'Система перенаправляет на главную страницу. '
            'В правом верхнем углу отображается имя пользователя.'
        ),
        (
            'TC-003',
            'Просмотр списка достопримечательностей по категории',
            'Пользователь авторизован. Находится на главной странице.',
            '1. Нажать на категорию достопримечательностей в меню\n'
            '2. Просмотреть отфильтрованный список\n'
            '3. Нажать на карточку достопримечательности',
            'Отображается список достопримечательностей выбранной категории. '
            'При нажатии на карточку открывается страница с детальной '
            'информацией о достопримечательности (фото, описание, рейтинг).'
        ),
        (
            'TC-004',
            'Бронирование билета на достопримечательность',
            'Пользователь авторизован. На странице деталей '
            'достопримечательности отображается информация о билетах.',
            '1. Выбрать тип билета\n'
            '2. Указать количество билетов (1)\n'
            '3. Нажать кнопку «Оформить заказ»\n'
            '4. Проверить корректность данных заказа\n'
            '5. Нажать «Перейти к оплате»',
            'Создан заказ со статусом «Ожидает оплаты». '
            'Система перенаправляет на страницу оплаты. '
            'В личном кабинете в разделе «Мои заказы» отображается новый заказ.'
        ),
        (
            'TC-005',
            'Успешная оплата заказа через AliPay',
            'Пользователь авторизован. Заказ создан. '
            'Открыта страница оплаты.',
            '1. Нажать кнопку «Оплатить через AliPay»\n'
            '2. Подтвердить оплату в мок-форме AliPay\n'
            '3. Дождаться перенаправления на страницу результата',
            'Отображается страница «Оплата успешна». '
            'Статус заказа изменён на «Оплачен». '
            'В истории заказов отображается оплаченный заказ.'
        ),
        (
            'TC-006',
            'Добавление комментария к путеводителю',
            'Пользователь авторизован. Открыта страница деталей путеводителя.',
            '1. Прокрутить вниз до блока комментариев\n'
            '2. Ввести текст комментария\n'
            '3. Нажать кнопку «Отправить»',
            'Комментарий опубликован под путеводителем. '
            'Отображается имя автора, текст комментария и время публикации.'
        ),
        (
            'TC-007',
            'Управление достопримечательностью (администратор)',
            'Администратор авторизован. '
            'Открыта панель управления достопримечательностями.',
            '1. Нажать кнопку «Добавить достопримечательность»\n'
            '2. Заполнить все поля формы (название, описание, категория, фото)\n'
            '3. Нажать «Сохранить»\n'
            '4. Найти новую достопримечательность в списке\n'
            '5. Нажать кнопку «Редактировать»\n'
            '6. Изменить название и нажать «Сохранить»\n'
            '7. Нажать кнопку «Удалить» и подтвердить удаление',
            'Новая достопримечательность создана и отображается в списке. '
            'После редактирования название обновлено. '
            'После удаления достопримечательность исчезает из списка.'
        ),
    ]

    for tc_id, title, preconds, steps, expected in test_cases:
        add_test_case_table(doc, tc_id, title, preconds, steps, expected)
        blank(doc, 1)

    out = os.path.join(os.path.dirname(__file__),
                       'ПР2_Тест_кейсы_Чэнь_Гэ.docx')
    doc.save(out)
    print(f'✓ Сохранён: {out}')
    return out

if __name__ == '__main__':
    create()
