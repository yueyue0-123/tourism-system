"""
fix_table_position.py — Перемещает таблицу нагрузочного тестирования в правильное место
"""

from docx import Document
from lxml import etree

DOC = r'C:\Users\大猪鼻\Desktop\godss\ПЗ_Образец_Чэнь_Гэ_5Этап.docx'
doc = Document(DOC)

# Найти таблицу нагрузочного тестирования (6 строк, 3 колонки, заголовки Эндпоинт/Среднее/Макс)
target_table = None
for t in doc.tables:
    if len(t.rows) == 6 and len(t.columns) == 3:
        first_cell_text = t.rows[0].cells[0].text.strip()
        if first_cell_text == 'Эндпоинт':
            target_table = t
            break

if not target_table:
    print('❌ Таблица не найдена')
    exit(1)

# Найти параграф "Результаты нагрузочного тестирования подтвердили..."
target_para = None
for p in doc.paragraphs:
    if 'Результаты нагрузочного тестирования подтвердили' in p.text:
        target_para = p
        break

if not target_para:
    print('❌ Параграф не найден')
    exit(1)

# Переместить таблицу перед параграфом с результатами
tbl_element = target_table._element
tbl_element.getparent().remove(tbl_element)
target_para._element.addprevious(tbl_element)

doc.save(DOC)
print('✅ Таблица перемещена перед параграфом с результатами')
