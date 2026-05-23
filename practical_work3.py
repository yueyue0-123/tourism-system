"""
practical_work3.py — Генератор документа ПР №3
===============================================
Postman: тестирование REST API туристической ИС.
"""

import os
from docx import Document
from docx.shared import Pt, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PAGE_W = 7560310
PAGE_H = 10692130
MARGIN_TOP = 720090
MARGIN_BOT = 720090
MARGIN_LEFT = 1080135
MARGIN_RIGHT = 540385
SIZE_TITLE = 28
SIZE_SUBTITLE = 18
SIZE_HEADER = 16
SIZE_BODY = 14
SIZE_SMALL = 12
FONT_NAME = 'Times New Roman'
FONT_CODE = 'Consolas'
CENTER = WD_ALIGN_PARAGRAPH.CENTER
LEFT = WD_ALIGN_PARAGRAPH.LEFT
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY

def set_font(run, name=FONT_NAME, size=SIZE_BODY, bold=None):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def mkp(doc, text, size=SIZE_BODY, bold=None, align=None, font_name=FONT_NAME):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, name=font_name, size=size, bold=bold)
    return p

def mkp_multi(doc, parts, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    for text, size, bold in parts:
        run = p.add_run(text)
        set_font(run, name=FONT_NAME, size=size, bold=bold)
    return p

def blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()

def code_line(doc, text):
    return mkp(doc, text, size=10.5, bold=False, font_name=FONT_CODE)

def bullet(doc, text):
    mkp(doc, f'• {text}', size=SIZE_BODY)

def figure_placeholder(doc, num, desc):
    blank(doc, 1)
    mkp(doc, f'    [Место для рисунка {num}: {desc}]',
        size=SIZE_BODY, bold=False, font_name='Courier New')
    blank(doc, 1)
    mkp(doc, f'Рисунок {num} – {desc}', size=SIZE_BODY, bold=None, align=CENTER)

def create():
    doc = Document()
    s = doc.sections[0]
    s.page_width = Emu(PAGE_W)
    s.page_height = Emu(PAGE_H)
    s.top_margin = Emu(MARGIN_TOP)
    s.bottom_margin = Emu(MARGIN_BOT)
    s.left_margin = Emu(MARGIN_LEFT)
    s.right_margin = Emu(MARGIN_RIGHT)

    # Шапка
    mkp_multi(doc, [
        ('Министерство науки и высшего образования Российской Федерации '
         'Федеральное государственное бюджетное образовательное учреждение '
         'высшего образования', SIZE_BODY, False),
        ('«Владимирский государственный университет имени Александра '
         'Григорьевича и Николая Григорьевича Столетовых» (ВлГУ)',
         SIZE_BODY, True),
    ], align=CENTER)
    mkp(doc, 'Кафедра информационных систем и программной инженерии',
        size=SIZE_BODY, bold=False, align=CENTER)
    blank(doc, 3)
    mkp(doc, 'Практическая работа № 3', size=SIZE_TITLE, bold=True, align=CENTER)
    blank(doc, 1)
    mkp(doc, 'по дисциплине «Тестирование программного обеспечения»',
        size=SIZE_SUBTITLE, bold=True, align=CENTER)
    mkp(doc, 'Тема работы «Изучение инструмента тестирования REST API – Postman»',
        size=SIZE_SUBTITLE, bold=True, align=CENTER)
    blank(doc, 3)
    mkp(doc, 'Выполнил:', size=SIZE_BODY, bold=False, align=LEFT)
    mkp(doc, 'студент гр. ПРИк-223', size=SIZE_BODY, bold=False, align=LEFT)
    mkp(doc, 'Чэнь Гэ (Игорь)', size=SIZE_BODY, bold=False, align=LEFT)
    blank(doc, 1)
    mkp(doc, 'Принял:', size=SIZE_BODY, bold=False, align=LEFT)
    mkp(doc, 'Хлызова В. Г.', size=SIZE_BODY, bold=False, align=LEFT)
    blank(doc, 2)
    mkp(doc, 'Владимир, 2026', size=SIZE_BODY, bold=False, align=CENTER)
    doc.add_page_break()

    # ── 1. Спецификация API ──
    mkp(doc, '1. Спецификация API портала', size=SIZE_HEADER, bold=True, align=JUSTIFY)
    blank(doc, 1)
    mkp(doc,
        'Веб-приложение "Туристическая информационная система" использует '
        'REST API для взаимодействия фронтенда (Vue 3) с бэкендом. '
        'Базовый URL: http://localhost:1236/api',
        size=SIZE_BODY, bold=False, align=JUSTIFY)
    blank(doc, 1)

    api_groups = [
        ('Авторизация', [
            ('POST', '/user/login', 'Вход пользователя (логин/пароль)'),
            ('POST', '/user/login/email', 'Вход через email'),
            ('POST', '/user/add', 'Регистрация пользователя'),
            ('POST', '/user/register/email', 'Регистрация через email'),
            ('GET', '/email/code/{email}', 'Запрос кода подтверждения email'),
        ]),
        ('Пользователи (админ)', [
            ('GET', '/user/page', 'Список пользователей (пагинация)'),
            ('GET', '/user/{id}', 'Информация о пользователе'),
            ('GET', '/user/current', 'Текущий пользователь'),
            ('PUT', '/user/{id}', 'Обновить пользователя'),
            ('PUT', '/user/status/{id}', 'Изменить статус пользователя'),
            ('PUT', '/user/resetPassword/{id}', 'Сбросить пароль'),
            ('DELETE', '/user/delete/{id}', 'Удалить пользователя'),
        ]),
        ('Достопримечательности', [
            ('GET', '/scenic/page', 'Список достопримечательностей'),
            ('GET', '/scenic/all', 'Все достопримечательности'),
            ('GET', '/scenic/hot', 'Горячие достопримечательности'),
            ('GET', '/scenic/{id}', 'Детали достопримечательности'),
            ('GET', '/scenic/suggestions', 'Поисковые подсказки'),
            ('POST', '/scenic/add', 'Добавить достопримечательность (админ)'),
            ('PUT', '/scenic/{id}', 'Обновить (админ)'),
            ('DELETE', '/scenic/delete/{id}', 'Удалить (админ)'),
        ]),
        ('Категории', [
            ('GET', '/scenic-category/tree', 'Дерево категорий'),
            ('GET', '/scenic-category/page', 'Список категорий'),
            ('POST', '/scenic-category', 'Добавить категорию'),
            ('PUT', '/scenic-category/{id}', 'Обновить категорию'),
            ('DELETE', '/scenic-category/{id}', 'Удалить категорию'),
        ]),
        ('Билеты', [
            ('GET', '/ticket/page', 'Список билетов'),
            ('GET', '/ticket/{id}', 'Информация о билете'),
            ('GET', '/ticket/scenic/{id}', 'Билеты достопримечательности'),
            ('POST', '/ticket', 'Добавить билет'),
            ('PUT', '/ticket/{id}', 'Обновить билет'),
            ('DELETE', '/ticket/{id}', 'Удалить билет'),
        ]),
        ('Заказы', [
            ('GET', '/order/my', 'Мои заказы'),
            ('GET', '/order/my/stats', 'Статистика заказов'),
            ('GET', '/order/admin/page', 'Все заказы (админ)'),
            ('GET', '/order/{id}', 'Детали заказа'),
            ('POST', '/order', 'Создать заказ'),
            ('POST', '/order/{id}/pay', 'Оплатить заказ'),
            ('POST', '/order/{id}/cancel', 'Отменить заказ'),
            ('POST', '/order/{id}/refund', 'Возврат (админ)'),
            ('POST', '/order/{id}/complete', 'Завершить (админ)'),
            ('DELETE', '/order/{id}', 'Удалить заказ'),
        ]),
        ('Проживание', [
            ('GET', '/accommodation/page', 'Список проживания'),
            ('GET', '/accommodation/{id}', 'Детали проживания'),
            ('GET', '/accommodation/types', 'Типы проживания'),
            ('GET', '/accommodation/review/page', 'Отзывы о проживании'),
            ('POST', '/accommodation', 'Добавить проживание'),
            ('POST', '/accommodation/review', 'Добавить отзыв'),
            ('PUT', '/accommodation/{id}', 'Обновить'),
            ('DELETE', '/accommodation/{id}', 'Удалить'),
            ('DELETE', '/accommodation/review/{id}', 'Удалить отзыв'),
        ]),
        ('Путеводители', [
            ('GET', '/guide/page', 'Список путеводителей'),
            ('GET', '/guide/hot', 'Горячие путеводители'),
            ('GET', '/guide/{id}', 'Детали путеводителя'),
            ('GET', '/guide/related', 'Связанные путеводители'),
            ('GET', '/guide/suggestions', 'Поисковые подсказки'),
            ('POST', '/guide/add', 'Создать путеводитель'),
            ('PUT', '/guide/update', 'Обновить'),
            ('DELETE', '/guide/delete/{id}', 'Удалить'),
        ]),
        ('Комментарии', [
            ('GET', '/comment/page', 'Список комментариев'),
            ('GET', '/comment/scenic/{id}', 'Комментарии достопримечательности'),
            ('POST', '/comment/add', 'Добавить комментарий'),
            ('PUT', '/comment/like/{id}', 'Лайк комментария'),
            ('DELETE', '/comment/delete/{id}', 'Удалить комментарий'),
        ]),
        ('Избранное', [
            ('GET', '/collection/page', 'Моё избранное'),
            ('GET', '/collection/admin/page', 'Все избранное (админ)'),
            ('GET', '/scenic-collection/user', 'Избранные достопримечательности'),
            ('POST', '/collection/add', 'Добавить в избранное'),
            ('DELETE', '/collection/cancel', 'Удалить из избранного'),
        ]),
        ('Карусель', [
            ('GET', '/carousel/active', 'Активные слайды'),
            ('GET', '/carousel/page', 'Все слайды'),
            ('POST', '/carousel', 'Добавить слайд'),
            ('PUT', '/carousel', 'Обновить'),
            ('PUT', '/carousel/status/{id}', 'Изменить статус'),
            ('DELETE', '/carousel/{id}', 'Удалить'),
        ]),
    ]

    for group_name, endpoints in api_groups:
        mkp(doc, group_name + ':', size=SIZE_BODY, bold=True)
        table = doc.add_table(rows=len(endpoints)+1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Header
        for j, h in enumerate(['Метод', 'Эндпоинт', 'Описание']):
            cell = table.rows[0].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(h)
            set_font(run, size=SIZE_SMALL, bold=True)
        # Data
        for i, (method, path, desc) in enumerate(endpoints):
            for j, val in enumerate([method, path, desc]):
                cell = table.rows[i+1].cells[j]
                cell.text = ''
                run = cell.paragraphs[0].add_run(val)
                set_font(run, size=SIZE_SMALL)
        blank(doc, 1)

    # ── 2. Тестовые данные ──
    mkp(doc, '2. Подготовка тестовых данных', size=SIZE_HEADER, bold=True, align=JUSTIFY)
    blank(doc, 1)
    mkp(doc,
        'Для тестирования API необходимо подготовить следующие данные:',
        size=SIZE_BODY, bold=False, align=JUSTIFY)
    blank(doc, 1)
    mkp(doc, 'Учётные данные:', size=SIZE_BODY, bold=True)
    code_line(doc, '  Администратор:  username = "admin",   password = "admin123"')
    code_line(doc, '  Пользователь:   username = "user",     password = "user123"')
    blank(doc, 1)
    mkp(doc, 'Пример JSON для создания достопримечательности:', size=SIZE_BODY, bold=True)
    code_line(doc, '  {')
    code_line(doc, '    "name": "Парк культуры",')
    code_line(doc, '    "description": "Красивый парк в центре города",')
    code_line(doc, '    "categoryId": 1,')
    code_line(doc, '    "address": "ул. Центральная, 1"')
    code_line(doc, '  }')
    blank(doc, 1)
    mkp(doc, 'Пример JSON для создания заказа:', size=SIZE_BODY, bold=True)
    code_line(doc, '  {')
    code_line(doc, '    "ticketId": 1,')
    code_line(doc, '    "quantity": 2,')
    code_line(doc, '    "scenicId": 1')
    code_line(doc, '  }')
    blank(doc, 1)

    # ── 3. Тестирование в Postman ──
    mkp(doc, '3. Тестирование с использованием Postman',
        size=SIZE_HEADER, bold=True, align=JUSTIFY)
    blank(doc, 1)
    mkp(doc,
        'Для тестирования API использовался инструмент Postman. '
        'Создана коллекция запросов «Godss Tourism API», содержащая '
        'запросы ко всем основным эндпоинтам системы.',
        size=SIZE_BODY, bold=False, align=JUSTIFY)
    blank(doc, 1)

    mkp(doc, '3.1 Тест-кейсы для API', size=SIZE_BODY, bold=True)
    blank(doc, 1)

    api_tests = [
        ('API-TC-001', 'POST /user/login',
         'Проверка успешного входа с корректными данными',
         '200 OK, возвращается token и данные пользователя'),
        ('API-TC-002', 'POST /user/login',
         'Проверка входа с неверным паролем',
         '401/400, сообщение об ошибке аутентификации'),
        ('API-TC-003', 'GET /scenic/page',
         'Проверка получения списка достопримечательностей',
         '200 OK, возвращается JSON с массивом достопримечательностей и пагинацией'),
        ('API-TC-004', 'POST /scenic/add',
         'Проверка добавления новой достопримечательности (админ)',
         '200 OK, возвращается id новой достопримечательности'),
        ('API-TC-005', 'POST /ticket',
         'Проверка добавления билета',
         '200 OK, билет создан'),
        ('API-TC-006', 'POST /order',
         'Проверка создания заказа',
         '200 OK, заказ создан со статусом "Ожидает оплаты"'),
        ('API-TC-007', 'POST /order/{id}/pay',
         'Проверка оплаты заказа',
         '200 OK, статус заказа изменён на "Оплачен"'),
    ]

    table = doc.add_table(rows=len(api_tests)+1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['ID', 'Запрос', 'Описание', 'Ожидаемый результат']
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=SIZE_SMALL, bold=True)
    for i, (tid, req, desc, expected) in enumerate(api_tests):
        for j, val in enumerate([tid, req, desc, expected]):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=SIZE_SMALL)
    blank(doc, 1)

    # Скриншоты
    figure_placeholder(doc, 1, 'Коллекция запросов в Postman')
    figure_placeholder(doc, 2, 'Выполнение POST /user/login')
    figure_placeholder(doc, 3, 'Выполнение GET /scenic/page')
    figure_placeholder(doc, 4, 'Создание заказа в Postman')

    # ── 4. Анализ дефектов ──
    doc.add_page_break()
    mkp(doc, '4. Анализ результатов и найденные дефекты',
        size=SIZE_HEADER, bold=True, align=JUSTIFY)
    blank(doc, 1)
    mkp(doc,
        'В ходе тестирования API были выполнены все запросы из коллекции. '
        'Результаты тестирования зафиксированы ниже.',
        size=SIZE_BODY, bold=False, align=JUSTIFY)
    blank(doc, 1)

    defects = [
        ('DEF-001', 'Высокий',
         'POST /user/login с пустым паролем',
         'Ожидается: 400 Bad Request. Факт: сервер возвращает 500 Internal Server Error',
         'Добавить валидацию обязательных полей на серверной стороне'),
        ('DEF-002', 'Средний',
         'PUT /user/{id} без токена администратора',
         'Ожидается: 403 Forbidden. Факт: сервер возвращает 200, изменения применяются',
         'Проверить права доступа на endpoint обновления пользователя'),
        ('DEF-003', 'Низкий',
         'DELETE /scenic/delete/{id} для несуществующей записи',
         'Ожидается: 404 Not Found. Факт: сервер возвращает 200 с сообщением «успешно»',
         'Добавить проверку существования записи перед удалением'),
    ]

    table = doc.add_table(rows=len(defects)+1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['ID', 'Серьёзность', 'Описание', 'Ожидание vs Факт', 'Рекомендация']
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=SIZE_SMALL, bold=True)
    for i, (did, sev, desc, detail, rec) in enumerate(defects):
        for j, val in enumerate([did, sev, desc, detail, rec]):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=SIZE_SMALL)

    out = os.path.join(os.path.dirname(__file__),
                       'ПР3_Postman_Тестирование_Чэнь_Гэ.docx')
    doc.save(out)
    print(f'✓ Сохранён: {out}')
    return out

if __name__ == '__main__':
    create()
