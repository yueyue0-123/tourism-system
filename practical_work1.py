"""
practical_work1.py — Генератор документа ПР №1
===============================================
Сценарии использования и чек-листы для туристической ИС.
"""

import os
from docx import Document
from docx.shared import Pt, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Параметры (как в шаблоне) ──
PAGE_W = 7560310
PAGE_H = 10692130
MARGIN_TOP = 720090
MARGIN_BOT = 720090
MARGIN_LEFT = 1080135
MARGIN_RIGHT = 540385
SIZE_TITLE = 28
SIZE_SUBTITLE = 18
SIZE_HEADER = 16
SIZE_BODY = 14
FONT_NAME = 'Times New Roman'
CENTER = WD_ALIGN_PARAGRAPH.CENTER
LEFT = WD_ALIGN_PARAGRAPH.LEFT
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY

def set_font(run, name=FONT_NAME, size=SIZE_BODY, bold=None):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def mkp(doc, text, size=SIZE_BODY, bold=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, name=FONT_NAME, size=size, bold=bold)
    return p

def mkp_multi(doc, parts, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    for text, size, bold in parts:
        run = p.add_run(text)
        set_font(run, name=FONT_NAME, size=size, bold=bold)
    return p

def blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()

def bullet(doc, text, indent=0):
    p = doc.add_paragraph()
    prefix = '    ' * indent + '• '
    run = p.add_run(prefix + text)
    set_font(run, size=SIZE_BODY)
    return p

def create():
    doc = Document()
    s = doc.sections[0]
    s.page_width = Emu(PAGE_W)
    s.page_height = Emu(PAGE_H)
    s.top_margin = Emu(MARGIN_TOP)
    s.bottom_margin = Emu(MARGIN_BOT)
    s.left_margin = Emu(MARGIN_LEFT)
    s.right_margin = Emu(MARGIN_RIGHT)

    # Шапка
    mkp_multi(doc, [
        ('Министерство науки и высшего образования Российской Федерации '
         'Федеральное государственное бюджетное образовательное учреждение '
         'высшего образования',
         SIZE_BODY, False),
        ('«Владимирский государственный университет имени Александра '
         'Григорьевича и Николая Григорьевича Столетовых» (ВлГУ)',
         SIZE_BODY, True),
    ], align=CENTER)
    mkp(doc, 'Кафедра информационных систем и программной инженерии',
        size=SIZE_BODY, bold=False, align=CENTER)
    blank(doc, 3)

    mkp(doc, 'Практическая работа № 1', size=SIZE_TITLE, bold=True, align=CENTER)
    blank(doc, 1)
    mkp(doc, 'по дисциплине «Тестирование программного обеспечения»',
        size=SIZE_SUBTITLE, bold=True, align=CENTER)
    mkp(doc,
        'Тема работы «Сценарии использования. Чек-листы»',
        size=SIZE_SUBTITLE, bold=True, align=CENTER)
    blank(doc, 3)

    mkp(doc, 'Выполнил:', size=SIZE_BODY, bold=False, align=LEFT)
    mkp(doc, 'студент гр. ПРИк-223', size=SIZE_BODY, bold=False, align=LEFT)
    mkp(doc, 'Чэнь Гэ (Игорь)', size=SIZE_BODY, bold=False, align=LEFT)
    blank(doc, 1)
    mkp(doc, 'Принял:', size=SIZE_BODY, bold=False, align=LEFT)
    mkp(doc, 'Хлызова В. Г.', size=SIZE_BODY, bold=False, align=LEFT)
    blank(doc, 2)
    mkp(doc, 'Владимир, 2026', size=SIZE_BODY, bold=False, align=CENTER)
    doc.add_page_break()

    # ── 1. Описание предметной области ──
    mkp(doc, '1. Описание предметной области', size=SIZE_HEADER, bold=True, align=JUSTIFY)
    blank(doc, 1)
    mkp(doc,
        'Разрабатываемая информационная система представляет собой '
        'веб-приложение для туристической сферы — "Туристическая '
        'информационная система" (Туристическая ИС). Система '
        'предназначена для предоставления пользователям информации '
        'о достопримечательностях, возможности бронирования билетов '
        'и проживания, а также создания и просмотра путеводителей.',
        size=SIZE_BODY, bold=False, align=JUSTIFY)

    mkp(doc,
        'Основные функциональные возможности системы:',
        size=SIZE_BODY, bold=False, align=JUSTIFY)
    for f in [
        'Регистрация и авторизация пользователей (по email и логину)',
        'Просмотр и поиск достопримечательностей по категориям',
        'Просмотр детальной информации о достопримечательностях',
        'Бронирование и покупка билетов',
        'Бронирование проживания рядом с достопримечательностями',
        'Создание и просмотр путеводителей (гидов)',
        'Система комментариев и оценок',
        'Добавление в избранное (коллекции)',
        'Онлайн-оплата через AliPay',
        'Личный кабинет пользователя',
        'Административная панель управления всеми разделами',
    ]:
        bullet(doc, f)
    blank(doc, 1)

    # ── 2. Сценарии использования ──
    mkp(doc, '2. Сценарии использования', size=SIZE_HEADER, bold=True, align=JUSTIFY)
    blank(doc, 1)

    mkp(doc, 'Название: «Покупка билета на достопримечательность»',
        size=SIZE_BODY, bold=True, align=JUSTIFY)
    mkp(doc,
        'Предусловие: Пользователь зарегистрирован, авторизован в системе '
        'и находится на главной странице.',
        size=SIZE_BODY, bold=False, align=JUSTIFY)
    blank(doc, 1)

    mkp(doc, 'Основной сценарий:', size=SIZE_BODY, bold=True, align=JUSTIFY)
    steps_main = [
        '1. Пользователь просматривает список достопримечательностей на главной странице.',
        '2. Пользователь выбирает достопримечательность и переходит на страницу с детальной информацией.',
        '3. Система отображает подробную информацию: описание, фотографии, категорию, рейтинг.',
        '4. Пользователь нажимает кнопку «Купить билет» и переходит на страницу бронирования.',
        '5. Система отображает доступные типы билетов и их стоимость.',
        '6. Пользователь выбирает количество билетов и нажимает «Оформить заказ».',
        '7. Система создает заказ и перенаправляет пользователя на страницу оплаты.',
        '8. Пользователь подтверждает оплату через AliPay.',
        '9. Система подтверждает оплату, обновляет статус заказа на «Оплачен».',
        '10. Система отправляет подтверждение на email пользователя.',
    ]
    for step in steps_main:
        bullet(doc, step)
    blank(doc, 1)

    mkp(doc, 'Альтернативный сценарий (отмена заказа):',
        size=SIZE_BODY, bold=True, align=JUSTIFY)
    steps_alt = [
        '1–6. Совпадают с основным сценарием.',
        '7. Пользователь передумал и нажимает кнопку «Отменить заказ» на странице оплаты.',
        '8. Система отменяет заказ, статус заказа меняется на «Отменён».',
        '9. Система перенаправляет пользователя на страницу «Мои заказы».',
    ]
    for step in steps_alt:
        bullet(doc, step)
    blank(doc, 2)

    # ── 3. Чек-лист ──
    mkp(doc, '3. Чек-лист для тестирования функционала приложения',
        size=SIZE_HEADER, bold=True, align=JUSTIFY)
    blank(doc, 1)

    checklist = [
        ('Регистрация и авторизация', [
            'Регистрация нового пользователя через email',
            'Регистрация нового пользователя через логин',
            'Вход в систему с корректными данными',
            'Вход с неверным паролем',
            'Вход с пустыми полями',
            'Выход из системы',
        ]),
        ('Поиск и просмотр достопримечательностей', [
            'Отображение списка достопримечательностей на главной',
            'Просмотр детальной информации о достопримечательности',
            'Фильтрация достопримечательностей по категориям',
            'Поиск достопримечательностей через поисковую строку',
            'Отображение горячих/популярных достопримечательностей',
        ]),
        ('Бронирование билетов', [
            'Отображение списка доступных билетов',
            'Выбор количества билетов',
            'Оформление заказа',
            'Отмена заказа до оплаты',
            'Просмотр истории заказов',
        ]),
        ('Оплата', [
            'Переход на страницу оплаты',
            'Успешная оплата через AliPay',
            'Отображение результата оплаты',
            'Обновление статуса заказа после оплаты',
        ]),
        ('Проживание', [
            'Отображение списка доступного проживания',
            'Просмотр деталей проживания',
            'Фильтрация по типу проживания',
            'Просмотр отзывов о проживании',
            'Добавление отзыва о проживании',
        ]),
        ('Путеводители (Гиды)', [
            'Просмотр списка путеводителей',
            'Просмотр деталей путеводителя',
            'Создание нового путеводителя',
            'Редактирование своего путеводителя',
            'Удаление своего путеводителя',
            'Добавление путеводителя в избранное',
        ]),
        ('Комментарии', [
            'Добавление комментария к достопримечательности',
            'Добавление комментария к путеводителю',
            'Удаление своего комментария',
            'Лайк комментария',
        ]),
        ('Избранное', [
            'Добавление путеводителя в избранное',
            'Добавление достопримечательности в избранное',
            'Просмотр списка избранного',
            'Удаление из избранного',
        ]),
        ('Личный кабинет', [
            'Просмотр личной информации',
            'Редактирование профиля',
            'Смена пароля',
            'Просмотр своих заказов',
            'Просмотр своих путеводителей',
        ]),
        ('Административная панель', [
            'Вход в административную панель',
            'Просмотр дашборда',
            'Управление пользователями (просмотр, добавление, удаление)',
            'Управление достопримечательностями (CRUD)',
            'Управление категориями достопримечательностей',
            'Управление билетами',
            'Управление заказами (просмотр, возврат, завершение)',
            'Управление проживанием (CRUD)',
            'Управление путеводителями (просмотр, удаление)',
            'Управление комментариями (просмотр, удаление)',
            'Управление каруселью на главной странице',
        ]),
    ]

    for section_name, items in checklist:
        mkp(doc, section_name + ':', size=SIZE_BODY, bold=True, align=JUSTIFY)
        for item in items:
            bullet(doc, f'☐ {item}')
        blank(doc, 1)

    # Сохранение
    out = os.path.join(os.path.dirname(__file__),
                       'ПР1_Сценарии_Чеклист_Чэнь_Гэ.docx')
    doc.save(out)
    print(f'✓ Сохранён: {out}')
    return out

if __name__ == '__main__':
    create()
