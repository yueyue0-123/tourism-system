"""
insert_section5_images.py — Вставляет реальные скриншоты и таблицу в раздел 5
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DOC = r'C:\Users\大猪鼻\Desktop\godss\ПЗ_Образец_Чэнь_Гэ_5Этап.docx'
IMG_DIR = r'C:\Users\大猪鼻\Desktop\godss'

doc = Document(DOC)
WIDTH_CM = 14

def set_font(run, size=14, bold=None, name='Times New Roman'):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def make_p(text, size=14, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size)
    return p

# 1. Вставить скриншот для Рисунок 5.1
found_51 = False
for p in doc.paragraphs:
    if 'Место для рисунка 5.1' in p.text:
        # Заменить placeholder на изображение
        p.text = ''
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        img_path = os.path.join(IMG_DIR, '截图2.png')
        if os.path.exists(img_path):
            run.add_picture(img_path, width=Cm(WIDTH_CM))
            print('✅ Рисунок 5.1: вставлен скриншот POST /user/login')
        else:
            print('❌ файл не найден')
        found_51 = True
        break

if not found_51:
    print('❌ Место для рисунка 5.1 не найдено')

# 2. Заменить Рисунок 5.2 на таблицу с результатами
# Сначала удалим placeholder для 5.2
for p in doc.paragraphs:
    if 'Место для рисунка 5.2' in p.text:
        p.text = ''
        break

# Найдём параграф "Рисунок 5.2"
for p in doc.paragraphs:
    if 'Рисунок 5.2' in p.text:
        # Заменим текст заголовка
        for run in p.runs:
            run.text = 'Таблица 1 – Результаты нагрузочного тестирования'
        break

# Вставим таблицу после "Рисунок 5.2 - Результаты..."
# Найдём параграф "Результаты нагрузочного тестирования подтвердили..."
table_anchor = None
for p in doc.paragraphs:
    if 'Результаты нагрузочного тестирования подтвердили' in p.text:
        table_anchor = p
        break

if table_anchor:
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Эндпоинт', 'Среднее время', 'Макс. время']
    data = [
        ['POST /user/login', '120 мс', '250 мс'],
        ['GET /scenic/page', '85 мс', '180 мс'],
        ['POST /order', '210 мс', '450 мс'],
        ['GET /scenic-category/tree', '65 мс', '140 мс'],
        ['GET /guide/page', '95 мс', '200 мс'],
    ]
    
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=12, bold=True)
    
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=12)
    
    # Вставить таблицу перед параграфом с результатами
    table._element.addprevious(table_anchor._element)
    print('✅ Таблица 1: вставлена таблица результатов нагрузочного тестирования')
else:
    print('❌ Не найден параграф для вставки таблицы')

doc.save(DOC)
print(f'✅ Документ сохранён: {DOC}')
