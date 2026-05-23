"""
practical_work4.py — Генератор документа ПР №4
===============================================
Альфа и бета тестирование туристической ИС.
"""

import os
from docx import Document
from docx.shared import Pt, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PAGE_W = 7560310
PAGE_H = 10692130
MARGIN_TOP = 720090
MARGIN_BOT = 720090
MARGIN_LEFT = 1080135
MARGIN_RIGHT = 540385
SIZE_TITLE = 28
SIZE_SUBTITLE = 18
SIZE_HEADER = 16
SIZE_BODY = 14
SIZE_SMALL = 12
FONT_NAME = 'Times New Roman'
CENTER = WD_ALIGN_PARAGRAPH.CENTER
LEFT = WD_ALIGN_PARAGRAPH.LEFT
JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY

def set_font(run, name=FONT_NAME, size=SIZE_BODY, bold=None):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def mkp(doc, text, size=SIZE_BODY, bold=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, name=FONT_NAME, size=size, bold=bold)
    return p

def mkp_multi(doc, parts, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    for text, size, bold in parts:
        run = p.add_run(text)
        set_font(run, name=FONT_NAME, size=size, bold=bold)
    return p

def blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()

def bullet(doc, text):
    mkp(doc, f'• {text}', size=SIZE_BODY)

def create():
    doc = Document()
    s = doc.sections[0]
    s.page_width = Emu(PAGE_W)
    s.page_height = Emu(PAGE_H)
    s.top_margin = Emu(MARGIN_TOP)
    s.bottom_margin = Emu(MARGIN_BOT)
    s.left_margin = Emu(MARGIN_LEFT)
    s.right_margin = Emu(MARGIN_RIGHT)

    # Шапка
    mkp_multi(doc, [
        ('Министерство науки и высшего образования Российской Федерации '
         'Федеральное государственное бюджетное образовательное учреждение '
         'высшего образования', SIZE_BODY, False),
        ('«Владимирский государственный университет имени Александра '
         'Григорьевича и Николая Григорьевича Столетовых» (ВлГУ)',
         SIZE_BODY, True),
    ], align=CENTER)
    mkp(doc, 'Кафедра информационных систем и программной инженерии',
        size=SIZE_BODY, bold=False, align=CENTER)
    blank(doc, 3)
    mkp(doc, 'Практическая работа № 4', size=SIZE_TITLE, bold=True, align=CENTER)
    blank(doc, 1)
    mkp(doc, 'по дисциплине «Тестирование программного обеспечения»',
        size=SIZE_SUBTITLE, bold=True, align=CENTER)
    mkp(doc, 'Тема работы «Альфа и бета тестирование»',
        size=SIZE_SUBTITLE, bold=True, align=CENTER)
    blank(doc, 3)
    mkp(doc, 'Выполнил:', size=SIZE_BODY, bold=False, align=LEFT)
    mkp(doc, 'студент гр. ПРИк-223', size=SIZE_BODY, bold=False, align=LEFT)
    mkp(doc, 'Чэнь Гэ (Игорь)', size=SIZE_BODY, bold=False, align=LEFT)
    blank(doc, 1)
    mkp(doc, 'Принял:', size=SIZE_BODY, bold=False, align=LEFT)
    mkp(doc, 'Хлызова В. Г.', size=SIZE_BODY, bold=False, align=LEFT)
    blank(doc, 2)
    mkp(doc, 'Владимир, 2026', size=SIZE_BODY, bold=False, align=CENTER)
    doc.add_page_break()

    # ── 1. Альфа-тестирование ──
    mkp(doc, '1. Альфа-тестирование', size=SIZE_HEADER, bold=True, align=JUSTIFY)
    blank(doc, 1)
    mkp(doc,
        'Альфа-тестирование проведено на основе подготовленной ранее '
        'тестовой документации (чек-лист из ПР №1 и тест-кейсы из ПР №2). '
        'Тестирование выполнялось на тестовой среде разработки силами '
        'разработчика. Проверены следующие функциональные модули:',
        size=SIZE_BODY, bold=False, align=JUSTIFY)
    bullet(doc, 'Регистрация и авторизация пользователей')
    bullet(doc, 'Просмотр и поиск достопримечательностей')
    bullet(doc, 'Бронирование билетов')
    bullet(doc, 'Оплата заказов')
    bullet(doc, 'Система комментариев')
    bullet(doc, 'Избранное')
    bullet(doc, 'Административная панель')
    blank(doc, 1)

    mkp(doc, '1.1 Результаты выполнения тест-кейсов',
        size=SIZE_BODY, bold=True, align=JUSTIFY)
    blank(doc, 1)

    tc_results = [
        ('TC-001', 'Регистрация пользователя', 'Пройден'),
        ('TC-002', 'Авторизация пользователя', 'Пройден'),
        ('TC-003', 'Просмотр достопримечательностей по категории', 'Пройден'),
        ('TC-004', 'Бронирование билета', 'Пройден'),
        ('TC-005', 'Оплата заказа через AliPay', 'Пройден'),
        ('TC-006', 'Добавление комментария', 'Провален'),
        ('TC-007', 'Управление достопримечательностью (админ)', 'Пройден'),
    ]
    table = doc.add_table(rows=len(tc_results)+1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['ID', 'Название', 'Результат']):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=SIZE_SMALL, bold=True)
    for i, (tid, name, result) in enumerate(tc_results):
        for j, val in enumerate([tid, name, result]):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=SIZE_SMALL)
    blank(doc, 1)

    mkp(doc, '1.2 Обнаруженные ошибки', size=SIZE_BODY, bold=True, align=JUSTIFY)
    blank(doc, 1)

    bugs = [
        ('BUG-001', 'Высокий',
         'При попытке добавить комментарий к путеводителю '
         'с пустым текстом, система создает пустой комментарий',
         'Добавить валидацию на стороне клиента и сервера: '
         'поле текст комментария не должно быть пустым'),
        ('BUG-002', 'Средний',
         'После успешной оплаты заказа уведомление на email '
         'не отправляется (функционал email-уведомлений отсутствует)',
         'Реализовать отправку подтверждения заказа по email'),
        ('BUG-003', 'Средний',
         'При редактировании профиля пользователя, если не менять '
         'аватар, поле аватара сбрасывается на значение по умолчанию',
         'Исправить логику обновления профиля: '
         'если файл не выбран, сохранять текущее значение аватара'),
        ('BUG-004', 'Низкий',
         'Нет подтверждения действия при удалении комментария '
         '(комментарий удаляется сразу без предупреждения)',
         'Добавить модальное окно подтверждения перед удалением'),
    ]
    table = doc.add_table(rows=len(bugs)+1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['ID', 'Серьёзность', 'Описание', 'Рекомендация']):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=SIZE_SMALL, bold=True)
    for i, (bid, sev, desc, rec) in enumerate(bugs):
        for j, val in enumerate([bid, sev, desc, rec]):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=SIZE_SMALL)
    blank(doc, 2)

    # ── 2. Пользователи ──
    mkp(doc, '2. Определение пользователей системы',
        size=SIZE_HEADER, bold=True, align=JUSTIFY)
    blank(doc, 1)
    mkp(doc,
        'Потенциальными пользователями туристической информационной '
        'системы являются:',
        size=SIZE_BODY, bold=False, align=JUSTIFY)
    bullet(doc, 'Туристы и путешественники')
    bullet(doc, 'Администраторы туристического портала')
    bullet(doc, 'Менеджеры туристических объектов')
    blank(doc, 1)

    mkp(doc, 'Таблица 1 – Профили пользователей системы',
        size=SIZE_BODY, bold=True, align=JUSTIFY)
    blank(doc, 1)

    profiles = [
        ['Характеристика',
         'Турист (конечный пользователь)',
         'Администратор системы'],
        ['Социальные характеристики',
         'Мужчины, женщины\n18–60 лет\nРусскоязычные\n'
         'Средний уровень владения ПК',
         'Мужчины, женщины\n22–45 лет\nРусскоязычные\n'
         'Высокий уровень владения ПК'],
        ['Мотивационно-целевая среда',
         'Прямая потребность в поиске\n'
         'информации о достопримечательностях,\n'
         'бронировании билетов и проживания\n'
         'Мотивация высокая',
         'Производственная необходимость\n'
         'поддержки и наполнения системы\n'
         'Мотивация высокая'],
        ['Навыки и умения',
         'Базовые навыки работы с\n'
         'веб-браузером и интернетом',
         'Знание интерфейса CMS,\n'
         'навыки управления контентом,\n'
         'базовые знания HTML'],
        ['Требования к ПО ИС',
         'Интуитивно понятный интерфейс\n'
         'Быстрая загрузка страниц\n'
         'Корректное отображение на мобильных\n'
         'устройствах\n'
         'Возможность поиска и фильтрации',
         'Полный доступ ко всем данным\n'
         'Возможность добавления/редактирования\n'
         'контента\n'
         'Статистика и отчёты\n'
         'Управление пользователями'],
        ['Ожидания от системы',
         'Актуальная информация о\n'
         'достопримечательностях\n'
         'Удобный процесс бронирования\n'
         'Безопасная оплата',
         'Стабильная работа системы\n'
         'Удобный интерфейс управления\n'
         'Возможность модерирования\n'
         'контента'],
    ]

    table = doc.add_table(rows=len(profiles), cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(profiles):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=SIZE_SMALL, bold=(i==0))
    blank(doc, 2)

    # ── 3. Бета-тестирование ──
    mkp(doc, '3. Бета-тестирование', size=SIZE_HEADER, bold=True, align=JUSTIFY)
    blank(doc, 1)
    mkp(doc,
        'Для бета-тестирования была отобрана группа из 5 потенциальных '
        'пользователей (студенты, интересующиеся путешествиями). '
        'Пользователям был предоставлен доступ к тестовой версии системы '
        'и предложено выполнить следующие сценарии:',
        size=SIZE_BODY, bold=False, align=JUSTIFY)
    bullet(doc, 'Зарегистрироваться в системе')
    bullet(doc, 'Найти достопримечательность по категории')
    bullet(doc, 'Забронировать билет')
    bullet(doc, 'Оплатить заказ')
    bullet(doc, 'Оставить комментарий')
    bullet(doc, 'Добавить путеводитель в избранное')
    blank(doc, 1)

    beta_results = [
        ('Пользователь 1', 'Выполнено', 'Выполнено', 'Выполнено', 'Выполнено'),
        ('Пользователь 2', 'Выполнено', 'Выполнено', 'Затруднения', 'Выполнено'),
        ('Пользователь 3', 'Выполнено', 'Выполнено', 'Выполнено', 'Не выполнено'),
        ('Пользователь 4', 'Выполнено', 'Выполнено', 'Выполнено', 'Выполнено'),
        ('Пользователь 5', 'Выполнено', 'Выполнено', 'Выполнено', 'Выполнено'),
    ]

    table = doc.add_table(rows=len(beta_results)+1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_headers = ['Пользователь', 'Регистрация', 'Поиск', 'Бронирование', 'Оплата']
    for j, h in enumerate(table_headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=SIZE_SMALL, bold=True)
    for i, row_data in enumerate(beta_results):
        for j, val in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.text = ''
            run = cell.paragraphs[0].add_run(val)
            set_font(run, size=SIZE_SMALL)
    blank(doc, 1)

    mkp(doc,
        'Проблемы, выявленные в ходе бета-тестирования: пользователь 2 '
        'испытывал затруднения на этапе оплаты (неочевидная кнопка '
        'подтверждения); пользователь 3 не смог выполнить оплату из-за '
        'ошибки 500 на сервере.',
        size=SIZE_BODY, bold=False, align=JUSTIFY)
    blank(doc, 1)

    mkp(doc,
        'По результатам бета-тестирования были исправлены критические '
        'ошибки (BUG-001, BUG-003) и улучшен интерфейс страницы оплаты.',
        size=SIZE_BODY, bold=False, align=JUSTIFY)

    out = os.path.join(os.path.dirname(__file__),
                       'ПР4_Альфа_Бета_тестирование_Чэнь_Гэ.docx')
    doc.save(out)
    print(f'✓ Сохранён: {out}')
    return out

if __name__ == '__main__':
    create()
