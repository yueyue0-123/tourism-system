"""
add_section5_figures.py — Добавляет места для рисунков в раздел 5
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC = r'C:\Users\大猪鼻\Desktop\godss\ПЗ_Образец_Чэнь_Гэ_5Этап.docx'
doc = Document(DOC)

def make_para(doc, text, size=14, bold=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              font_name='Times New Roman'):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    return p

def insert_before(ref_para, new_para):
    ref_para._element.addprevious(new_para._element)

# Найдём конец раздела 5.2 и начало 5.3
target = None
for i, p in enumerate(doc.paragraphs):
    if '5.3 Нагрузочное' in p.text:
        target = p
        break

if target is None:
    print('❌ 5.3 не найдено')
    exit(1)

# Рисунок 5.1 – Postman API тестирование
f1_caption = make_para(doc, 'Рисунок 5.1 – Пример выполнения POST-запроса в Postman',
                       size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
f1_placeholder = make_para(doc, '    [Место для рисунка 5.1: скриншот Postman с запросом POST /user/login]',
                          size=10.5, font_name='Consolas', align=WD_ALIGN_PARAGRAPH.LEFT)
# Пустая строка перед
f1_blank = make_para(doc, '', size=14)

insert_before(target, f1_placeholder)
insert_before(target, f1_caption)
insert_before(target, f1_blank)

# Теперь найдём конец 5.3
target2 = None
for i, p in enumerate(doc.paragraphs):
    if 'Верификация работоспособности' in p.text:
        target2 = p
        break

if target2:
    f2_caption = make_para(doc, 'Рисунок 5.2 – Результаты нагрузочного тестирования в JMeter',
                          size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    f2_placeholder = make_para(doc, '    [Место для рисунка 5.2: скриншот Apache JMeter с результатами тестирования]',
                              size=10.5, font_name='Consolas', align=WD_ALIGN_PARAGRAPH.LEFT)
    f2_blank = make_para(doc, '', size=14)
    
    insert_before(target2, f2_placeholder)
    insert_before(target2, f2_caption)
    insert_before(target2, f2_blank)

doc.save(DOC)
print(f'✅ Рисунки 5.1 и 5.2 добавлены в документ')
