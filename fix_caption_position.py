"""
fix_caption_position.py — Перемещает подпись таблицы 1 в правильное место
"""

from docx import Document

DOC = r'C:\Users\大猪鼻\Desktop\godss\ПЗ_Образец_Чэнь_Гэ_5Этап.docx'
doc = Document(DOC)

# Найти параграф "Таблица 1 – Результаты..."
caption_p = None
for p in doc.paragraphs:
    if 'Таблица 1 – Результаты' in p.text:
        caption_p = p
        break

# Найти параграф результатов
result_p = None
for p in doc.paragraphs:
    if 'Результаты нагрузочного тестирования подтвердили' in p.text:
        result_p = p
        break

if caption_p and result_p:
    # Переместить caption перед results
    elem = caption_p._element
    elem.getparent().remove(elem)
    result_p._element.addprevious(elem)
    print('✅ Подпись таблицы перемещена перед параграфом с результатами')
else:
    print('❌ Не найдено')
    if not caption_p: print('  caption_p not found')
    if not result_p: print('  result_p not found')

# Также переместим саму таблицу, если она всё ещё не там
# Найти таблицу
for t in doc.tables:
    if len(t.rows) == 6 and len(t.columns) == 3:
        first = t.rows[0].cells[0].text.strip()
        if first == 'Эндпоинт':
            tbl_elem = t._element
            # Проверить, что таблица перед result_p (через сравнение XML siblings)
            # Если нет, переместить
            parent = tbl_elem.getparent()
            siblings = list(parent)
            tbl_idx = siblings.index(tbl_elem)
            result_idx = siblings.index(result_p._element)
            if tbl_idx > result_idx:
                parent.remove(tbl_elem)
                result_p._element.addprevious(tbl_elem)
                print('✅ Таблица перемещена')
            else:
                print('✅ Таблица уже на месте')
            break

doc.save(DOC)
print('✅ Готово')
