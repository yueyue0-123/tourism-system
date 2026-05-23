"""
generate_code_explanation.py — Документ с объяснением кода
=========================================================
Описание архитектуры и ключевых компонентов системы.
"""

import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'C:\Users\大猪鼻\Desktop\godss\Описание_кода_туристической_ИС.docx'

def set_font(run, size=14, bold=None, name='Times New Roman'):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def mkp(doc, text, size=14, bold=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p

def mkp_multi(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    for text, size, bold in parts:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold)
    return p

def blank(doc, n=1):
    for _ in range(n):
        doc.add_paragraph()

def bullet(doc, text):
    mkp(doc, f'• {text}')

def code_line(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, size=10.5, name='Consolas')

doc = Document()

# Титульная страница
mkp(doc, 'Министерство науки и высшего образования РФ', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
mkp(doc, 'ФГБОУ ВО "Владимирский государственный университет"', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
mkp(doc, 'Кафедра ИСПИ', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
blank(doc, 3)
mkp(doc, 'Описание программного кода', size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
mkp(doc, 'Распределённая информационная система путешествий', size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
blank(doc, 1)
mkp(doc, 'Дисциплина: "Распределённые программные системы"', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
blank(doc, 3)
mkp(doc, 'Выполнил: студент гр. ПРИк-223', size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
mkp(doc, 'Чэнь Гэ (Игорь)', size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
blank(doc, 1)
mkp(doc, 'Владимир, 2026', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# 1. Архитектура системы
mkp(doc, '1. Архитектура системы', size=16, bold=True)
blank(doc, 1)
mkp(doc, 'Система построена по классической трёхзвенной клиент-серверной архитектуре:')
bullet(doc, 'Клиентское приложение (Frontend) — Vue 3 + Element Plus')
bullet(doc, 'Сервер приложений (Backend) — Spring Boot 3.x на Java 17')
bullet(doc, 'База данных — MySQL с кэшированием Redis')
blank(doc, 1)
mkp(doc, 'Клиентское приложение взаимодействует с сервером через REST API, '
         'используя библиотеку Axios для HTTP-запросов. Серверная часть, в свою очередь, '
         'обеспечивает бизнес-логику, валидацию данных и доступ к базе данных через MyBatis-Plus.')
blank(doc, 1)

# 2. Структура проекта
mkp(doc, '2. Структура проекта (клиентская часть)', size=16, bold=True)
blank(doc, 1)
code_line(doc, 'godss/')
code_line(doc, '├── src/')
code_line(doc, '│   ├── api/           # API-вызовы (user.js)')
code_line(doc, '│   ├── assets/        # Статические ресурсы (CSS, изображения)')
code_line(doc, '│   ├── components/    # Переиспользуемые компоненты')
code_line(doc, '│   │   ├── backend/   #   Navbar, Sidebar')
code_line(doc, '│   │   ├── common/    #   SmartSearch')
code_line(doc, '│   │   ├── frontend/  #   CategoryMenu, HomeCarousel')
code_line(doc, '│   │   └── WangEditor.vue')
code_line(doc, '│   ├── layouts/       # Макеты страниц')
code_line(doc, '│   ├── router/        # Маршрутизация (index.js)')
code_line(doc, '│   ├── store/        # Состояние приложения (Pinia/Vuex)')
code_line(doc, '│   ├── styles/       # Глобальные стили')
code_line(doc, '│   ├── utils/        # Утилиты (request.js, dateUtils.js)')
code_line(doc, '│   └── views/        # Страницы приложения')
code_line(doc, '│       ├── auth/      #   Вход/регистрация')
code_line(doc, '│       ├── backend/   #   Админ-панель')
code_line(doc, '│       ├── frontend/  #   Пользовательский интерфейс')
code_line(doc, '│       └── error/     #   404')
code_line(doc, '├── vue.config.js     # Конфигурация Vue + прокси')
code_line(doc, '└── package.json      # Зависимости проекта')
blank(doc, 1)

# 3. Ключевые компоненты
mkp(doc, '3. Ключевые компоненты клиентской части', size=16, bold=True)
blank(doc, 1)

mkp(doc, '3.1 Конфигурация HTTP-запросов (src/utils/request.js)', size=14, bold=True)
blank(doc, 1)
mkp(doc, 'Файл request.js предоставляет единый интерфейс для выполнения HTTP-запросов. '
         'Он использует библиотеку Axios с предварительно настроенными перехватчиками:')
bullet(doc, 'Автоматическое добавление токена авторизации из localStorage')
bullet(doc, 'Единая обработка ошибок с отображением через Element Plus Message')
bullet(doc, 'Поддержка кастомных сообщений об успехе/ошибке')
bullet(doc, 'Callback-функции onSuccess/onError для гибкой обработки ответов')
blank(doc, 1)
code_line(doc, '// Пример использования:')
code_line(doc, 'request.get("/scenic/page", { page: 1, size: 10 })')
code_line(doc, 'request.post("/order", { ticketId: 1, quantity: 2 })')
code_line(doc, 'request.put("/user/1", { name: "New Name" })')
code_line(doc, 'request.delete("/scenic/delete/1")')
blank(doc, 1)

mkp(doc, '3.2 Маршрутизация (src/router/index.js)', size=14, bold=True)
blank(doc, 1)
mkp(doc, 'Маршрутизация реализована с использованием Vue Router. Система разделена на три группы маршрутов:')
bullet(doc, 'Фронтенд-маршруты (frontendRoutes) — главная страница, список достопримечательностей, '
         'детальная информация, бронирование, заказы, оплата, путеводители, избранное, профиль')
bullet(doc, 'Бэкенд-маршруты (backendRoutes) — дашборд, управление пользователями, '
         'достопримечательностями, категориями, комментариями, путеводителями, заказами и т.д.')
bullet(doc, 'Маршруты ошибок (errorRoutes) — страница 404')
blank(doc, 1)
mkp(doc, 'Реализована защита маршрутов через beforeEach guard:')
bullet(doc, 'Проверка аутентификации (requiresAuth)')
bullet(doc, 'Разделение доступа между обычными пользователями и администраторами')
bullet(doc, 'Перенаправление на страницу входа при отсутствии токена')
blank(doc, 1)

mkp(doc, '3.3 Главная страница (src/views/frontend/Home.vue)', size=14, bold=True)
blank(doc, 1)
mkp(doc, 'Главная страница содержит следующие секции:')
bullet(doc, 'Hero-секция с каруселью изображений и поисковой строкой')
bullet(doc, 'Быстрая навигация по категориям достопримечательностей')
bullet(doc, 'Избранные достопримечательности с карточками')
bullet(doc, 'Горячие путеводители')
blank(doc, 1)

mkp(doc, '3.4 Детальная страница достопримечательности (src/views/frontend/scenic/ScenicDetail.vue)', size=14, bold=True)
blank(doc, 1)
mkp(doc, 'Страница отображает полную информацию о достопримечательности:')
bullet(doc, 'Фотографии, описание, категория, рейтинг')
bullet(doc, 'Список доступных билетов с ценами')
bullet(doc, 'Блок комментариев с возможностью добавления')
bullet(doc, 'Кнопка добавления в избранное')
bullet(doc, 'Взаимодействие с сервером через API для получения данных')
blank(doc, 1)

mkp(doc, '3.5 Бронирование билетов (src/views/frontend/ticket/booking.vue)', size=14, bold=True)
blank(doc, 1)
mkp(doc, 'Компонент бронирования реализует следующий workflow:')
bullet(doc, 'Выбор типа и количества билетов')
bullet(doc, 'Отправка POST-запроса на /order для создания заказа')
bullet(doc, 'Перенаправление на оплату через AliPay')
bullet(doc, 'Возможность отмены заказа')
blank(doc, 1)

# 4. API эндпоинты
mkp(doc, '4. API эндпоинты', size=16, bold=True)
blank(doc, 1)
mkp(doc, 'Система предоставляет следующие группы REST API эндпоинтов:')
blank(doc, 1)
api_groups = [
    ('Авторизация', 'POST /user/login, POST /user/add, POST /user/login/email, GET /email/code/{email}'),
    ('Пользователи', 'GET /user/page, GET /user/{id}, PUT /user/{id}, DELETE /user/delete/{id}'),
    ('Достопримечательности', 'GET /scenic/page, GET /scenic/{id}, POST /scenic/add, PUT /scenic/{id}'),
    ('Категории', 'GET /scenic-category/tree, POST /scenic-category, PUT /scenic-category/{id}'),
    ('Билеты', 'GET /ticket/page, GET /ticket/{id}, POST /ticket, PUT /ticket/{id}'),
    ('Заказы', 'GET /order/my, POST /order, POST /order/{id}/pay, POST /order/{id}/cancel'),
    ('Путеводители', 'GET /guide/page, POST /guide/add, PUT /guide/update, DELETE /guide/delete/{id}'),
    ('Комментарии', 'GET /comment/page, POST /comment/add, PUT /comment/like/{id}'),
    ('Избранное', 'GET /collection/page, POST /collection/add, DELETE /collection/cancel'),
    ('Проживание', 'GET /accommodation/page, POST /accommodation, POST /accommodation/review'),
    ('Карусель', 'GET /carousel/active, POST /carousel, PUT /carousel'),
    ('Файлы', 'POST /file/upload/img'),
    ('Оплата', 'POST /alipay/mock-pay/{id}'),
]
for name, endpoints in api_groups:
    mkp(doc, f'• {name}:', size=14, bold=True)
    mkp(doc, f'  {endpoints}', size=12)
blank(doc, 1)

# 5. Конфигурация прокси
mkp(doc, '5. Конфигурация прокси (vue.config.js)', size=16, bold=True)
blank(doc, 1)
mkp(doc, 'Для обеспечения связи между клиентом и сервером используется прокси-сервер '
         'встроенного dev-сервера Vue CLI. Все запросы, начинающиеся с /api, '
         'перенаправляются на http://localhost:1236.')
blank(doc, 1)
code_line(doc, '// vue.config.js')
code_line(doc, "proxy: {")
code_line(doc, "  '/api': {")
code_line(doc, "    target: 'http://localhost:1236',")
code_line(doc, "    changeOrigin: true")
code_line(doc, "  }")
code_line(doc, "}")
blank(doc, 1)

# 6. Заключение
mkp(doc, '6. Заключение', size=16, bold=True)
blank(doc, 1)
mkp(doc, 'Разработанная распределённая информационная система путешествий представляет собой '
         'полнофункциональное веб-приложение, реализующее клиент-серверную архитектуру. '
         'Клиентская часть на Vue 3 обеспечивает современный адаптивный интерфейс с '
         'маршрутизацией, управлением состоянием и широким набором пользовательских сценариев. '
         'Серверная часть на Spring Boot предоставляет REST API с авторизацией, бизнес-логикой '
         'и доступом к данным через MyBatis-Plus. Система охватывает все ключевые функции '
         'туристического портала: от поиска достопримечательностей до онлайн-оплаты заказов.')

doc.save(OUT)
print(f'✅ Готово: {OUT}')
