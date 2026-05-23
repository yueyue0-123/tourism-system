"""
insert_pr3_screenshots.py — Вставляет скриншоты в ПР3
"""

import os
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC = r'C:\Users\大猪鼻\Desktop\godss\ПР3_Postman_Тестирование_Чэнь_Гэ.docx'
DIR = r'C:\Users\大猪鼻\Desktop\godss'

FIGURES = {
    1: '截图1.png',
    2: '截图2.png',
    3: '截图3.png',
    4: '截图4.png',
}

doc = Document(DOC)
WIDTH_CM = 14

inserted = 0
for p in doc.paragraphs:
    text = p.text.strip()
    if text.startswith('[Место для рисунка') and ']' in text:
        try:
            num_str = text.split('[')[1].split(']')[0]
            num_str = num_str.replace('Место для рисунка ', '')
            num = int(num_str.split(':')[0].strip())
        except:
            continue
        
        if num in FIGURES:
            img_path = os.path.join(DIR, FIGURES[num])
            if os.path.exists(img_path):
                p.text = ''
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Cm(WIDTH_CM))
                inserted += 1
                print(f'  ✅ Рисунок {num} вставлен')
            else:
                print(f'  ❌ Рисунок {num}: файл не найден')

print(f'\n✅ Вставлено {inserted} из {len(FIGURES)} скриншотов')
doc.save(DOC)
print(f'✅ Документ сохранён: {DOC}')
